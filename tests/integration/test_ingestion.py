"""PDF/DOCX 解析器测试。"""

from __future__ import annotations

import os
from pathlib import Path

import pytest

os.environ["MINERU_MODEL_SOURCE"] = "local"

FIXTURES_DIR = Path(__file__).parent.parent / "fixtures"
PDF_DIR = FIXTURES_DIR / "pdfs"


# ── Fixtures ──────────────────────────────────────────────


@pytest.fixture
def sample_content_list():
    """模拟结构化内容（兼容 MinerU 和 DOCX 两种输出格式）。"""
    return [
        {"type": "text", "text": "第一章 函数与极限", "text_level": 2, "page_idx": 0},
        {"type": "text", "text": "本节介绍映射与函数的概念。", "text_level": 99, "page_idx": 0},
        {"type": "text", "text": "一、映射", "text_level": 3, "page_idx": 0},
        {"type": "text", "text": "映射是现代数学的基本概念。", "text_level": 99, "page_idx": 0},
        {"type": "text", "text": "二、函数", "text_level": 3, "page_idx": 1},
        {"type": "text", "text": "函数是实数集到实数集的映射。", "text_level": 99, "page_idx": 1},
    ]


# ── parser_utils 单元测试 ──────────────────────────────


class TestSplitByHeadings:
    def test_basic_split(self, sample_content_list):
        from coursepilot.ingestion.parser_utils import _split_by_headings

        blocks = _split_by_headings(sample_content_list)
        assert len(blocks) == 3
        assert "第一章 函数与极限" in blocks[0]["text"]
        assert "一、映射" in blocks[1]["text"]
        assert "二、函数" in blocks[2]["text"]

    def test_page_ref(self, sample_content_list):
        from coursepilot.ingestion.parser_utils import _split_by_headings

        blocks = _split_by_headings(sample_content_list)
        print(blocks)
        assert "p2" in blocks[2]["page_ref"]  # "二、函数" 在 page_idx=1


class TestSplitTextV2:
    def test_short_text_no_split(self):
        from coursepilot.ingestion.parser_utils import _split_text_v2

        text = "短文本"
        chunks = _split_text_v2(text, target_chars=800, hard_lower=400, hard_upper=1200)
        assert chunks == [text]

    def test_long_text_is_split_with_paragraph_boundary(self):
        from coursepilot.ingestion.parser_utils import _split_text_v2

        # 两段各 ~600 字符，总长 ~1200 > hard_upper=900 → 应在段落边界切分
        para = "数学分析是研究函数的学科。" * 50  # ~600 chars
        text = para + "\n\n" + para
        chunks = _split_text_v2(text, target_chars=800, hard_lower=300, hard_upper=900)
        assert len(chunks) == 2
        assert chunks[0] == para

    def test_long_text_forced_split(self):
        from coursepilot.ingestion.parser_utils import _split_text_v2

        # 超长无段落边界的文本 → 在句边界或强制切分
        text = "这是测试文本。" * 200  # ~1400 chars
        chunks = _split_text_v2(text, target_chars=800, hard_lower=300, hard_upper=1000)
        assert len(chunks) >= 2
        for chunk in chunks:
            assert len(chunk) <= 1000  # 不超过 hard_upper

    def test_math_block_protected(self):
        from coursepilot.ingestion.parser_utils import _split_text_v2

        # 数学块 $$...$$ 跨越切分边界时，切分点应避开数学块
        before = "微积分基础知识。" * 30  # ~300 chars
        math_block = "$$\n\\int_a^b f(x)dx = F(b) - F(a)\n$$"
        after = "这是牛顿-莱布尼茨公式。" * 40  # ~500 chars
        text = before + "\n\n" + math_block + "\n\n" + after
        chunks = _split_text_v2(text, target_chars=800, hard_lower=300, hard_upper=1200)
        # 数学块应完整保留在某个 chunk 中
        math_in_chunks = [math_block in c for c in chunks]
        assert any(math_in_chunks), "数学块应完整保留"


class TestExtractKnowledgeUnits:
    def test_basic(self, sample_content_list):
        from coursepilot.ingestion.parser_utils import extract_knowledge_units

        doc_id = "00000000-0000-0000-0000-000000000001"
        kp_id = "00000000-0000-0000-0000-000000000002"
        units = extract_knowledge_units(
            sample_content_list,
            document_id=doc_id,
            kp_id=kp_id,
        )
        assert len(units) == 3
        for u in units:
            assert u["kp_id"] == kp_id
            assert u["document_id"] == doc_id
            assert u["content"] != ""

    def test_document_id_and_kp_id(self, sample_content_list):
        from coursepilot.ingestion.parser_utils import extract_knowledge_units

        units = extract_knowledge_units(
            sample_content_list,
            document_id="doc-1",
            kp_id="kp-1",
        )
        for u in units:
            assert u["document_id"] == "doc-1"
            assert u["kp_id"] == "kp-1"


# ── DOCX 解析器测试 ────────────────────────────────────


class TestDocxParser:
    def test_parse_docx_structure(self, tmp_path):
        """用 python-docx 创建一个临时 DOCX，验证结构提取。"""
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        doc.add_heading("第一章 函数与极限", level=1)
        doc.add_paragraph("本节介绍函数的基本概念。")
        doc.add_heading("第一节 映射", level=2)
        doc.add_paragraph("映射是现代数学的基本概念。")

        file_path = tmp_path / "test.docx"
        doc.save(str(file_path))

        from coursepilot.ingestion.docx_parser import parse_docx

        result = parse_docx(str(file_path))

        assert len(result["content_list"]) == 4
        assert result["content_list"][0]["text"] == "第一章 函数与极限"
        assert result["content_list"][0]["text_level"] == 1
        assert result["content_list"][2]["text"] == "第一节 映射"
        assert result["content_list"][2]["text_level"] == 2
        assert "第一章 函数与极限" in result["markdown"]

    def test_extract_knowledge_units_from_docx(self, tmp_path):
        """DOCX 解析后接 KnowledgeUnit 切片。"""
        from docx import Document

        doc = Document()
        doc.add_heading("第一章", level=1)
        doc.add_paragraph("内容一。")
        doc.add_heading("第二章", level=1)
        doc.add_paragraph("内容二。")

        file_path = tmp_path / "test2.docx"
        doc.save(str(file_path))

        from coursepilot.ingestion.docx_parser import parse_docx
        from coursepilot.ingestion.parser_utils import extract_knowledge_units

        parsed = parse_docx(str(file_path))
        units = extract_knowledge_units(
            parsed["content_list"],
            document_id="doc-1",
            kp_id="kp-1",
        )

        assert len(units) == 2
        assert "第一章" in units[0]["content"]
        assert "第二章" in units[1]["content"]


# ── PDF 集成测试（需要真实 PDF 文件 + MinerU）────────────


@pytest.mark.slow
@pytest.mark.asyncio
async def test_parse_pdf_small_file():
    output_dir = "output/report_test1"
    pdf = PDF_DIR / "report.pdf"
    if not pdf.exists():
        pytest.skip("report.pdf not found")

    from coursepilot.ingestion.pdf_parser import parse_pdf

    result = await parse_pdf(str(pdf), str(output_dir), start_page=0, end_page=1)

    assert "项目验收报告" in result["markdown"]
    assert len(result["content_list"]) > 0


@pytest.mark.slow
@pytest.mark.asyncio
async def test_parse_pdf_and_chunk():
    output_dir = "output/report_test2"
    pdf = PDF_DIR / "report.pdf"
    if not pdf.exists():
        pytest.skip("report.pdf not found")

    from coursepilot.ingestion.pdf_parser import parse_pdf
    from coursepilot.ingestion.parser_utils import extract_knowledge_units

    result = await parse_pdf(str(pdf), str(output_dir), start_page=0, end_page=1)
    print(result)
    units = extract_knowledge_units(
        result["content_list"],
        document_id="doc-1",
        kp_id="kp-1",
    )
    print(units)

    assert len(units) >= 1
    assert "项目验收报告" in units[0]["content"]
